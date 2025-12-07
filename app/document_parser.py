# app/document_parser.py
# Extract text from PDF and DOC/DOCX files

import io
import logging
from typing import Union
from pathlib import Path

# PDF parsing
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.pdfparser import PDFSyntaxError

# DOCX parsing
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file content, including attempting to find URLs.
    
    Args:
        file_content: PDF file as bytes
        
    Returns:
        Extracted text from PDF
        
    Raises:
        ValueError: If PDF cannot be parsed
    """
    try:
        text = pdf_extract_text(io.BytesIO(file_content))
        if not text or not text.strip():
            raise ValueError("No text content found in PDF")
        
        # Try to extract URLs from PDF annotations (if available)
        try:
            from pdfminer.pdfpage import PDFPage
            from pdfminer.pdfdocument import PDFDocument
            from pdfminer.pdfparser import PDFParser
            
            pdf_file = io.BytesIO(file_content)
            parser = PDFParser(pdf_file)
            document = PDFDocument(parser)
            
            urls = []
            for page in PDFPage.create_pages(document):
                if 'Annots' in page.attrs:
                    for annot in page.attrs['Annots']:
                        try:
                            annot_obj = annot.resolve()
                            if 'A' in annot_obj and 'URI' in annot_obj['A']:
                                url = annot_obj['A']['URI'].decode('utf-8') if isinstance(annot_obj['A']['URI'], bytes) else str(annot_obj['A']['URI'])
                                urls.append(url)
                        except:
                            pass
            
            # Append found URLs to the text
            if urls:
                text += "\n\nExtracted URLs:\n" + "\n".join(urls)
                logger.info(f"Extracted {len(urls)} URLs from PDF")
        except Exception as url_error:
            logger.debug(f"Could not extract URLs from PDF: {url_error}")
            # Continue with text-only extraction
        
        return text.strip()
    except PDFSyntaxError as e:
        logger.error(f"PDF syntax error: {e}")
        raise ValueError(f"Invalid PDF file: {e}")
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file content, including hyperlinks.
    
    Args:
        file_content: DOCX file as bytes
        
    Returns:
        Extracted text from DOCX with embedded URLs
        
    Raises:
        ValueError: If DOCX cannot be parsed
    """
    try:
        doc = Document(io.BytesIO(file_content))
        
        # Extract text from paragraphs with hyperlinks
        paragraphs = []
        for para in doc.paragraphs:
            para_text = []
            
            # Check for hyperlinks in the paragraph
            for run in para.runs:
                text = run.text.strip()
                if text:
                    # Try to find hyperlink
                    if hasattr(run, '_element') and run._element.xml:
                        # Check if this run contains a hyperlink
                        import re
                        hyperlink_match = re.search(r'<w:hyperlink[^>]*r:id=\"([^\"]*)\"', run._element.xml)
                        if hyperlink_match:
                            rel_id = hyperlink_match.group(1)
                            try:
                                # Get the actual URL from the relationship
                                rel = doc.part.rels[rel_id]
                                url = rel.target_ref
                                # Append text with URL
                                para_text.append(f"{text} ({url})")
                                continue
                            except:
                                pass
                    para_text.append(text)
            
            full_para = ' '.join(para_text).strip()
            if full_para:
                paragraphs.append(full_para)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        text = "\n".join(paragraphs)
        if not text or not text.strip():
            raise ValueError("No text content found in DOCX")
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_doc(file_content: bytes) -> str:
    """
    Extract text from old DOC format.
    Note: python-docx doesn't support old .doc format.
    This is a placeholder that returns an error message.
    For full .doc support, you'd need to use antiword, LibreOffice, or similar.
    
    Args:
        file_content: DOC file as bytes
        
    Returns:
        Error message
        
    Raises:
        ValueError: Old DOC format not supported
    """
    raise ValueError(
        "Old .doc format is not supported. Please save your resume as .docx or .pdf format."
    )


def extract_text_from_document(file_content: bytes, filename: str) -> str:
    """
    Extract text from a document file (PDF, DOC, or DOCX).
    
    Args:
        file_content: File content as bytes
        filename: Original filename to determine type
        
    Returns:
        Extracted text from the document
        
    Raises:
        ValueError: If file type is not supported or extraction fails
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_content)
    elif filename_lower.endswith('.doc'):
        return extract_text_from_doc(file_content)
    else:
        raise ValueError(
            f"Unsupported file type. Please upload a PDF or DOCX file. Got: {filename}"
        )
