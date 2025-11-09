# Generate the formatted resume
# create_resume(resume_data)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_paragraph_format(paragraph, font_name="Times New Roman", font_size=10, spacing=1.0, alignment=None):
    """Apply consistent font, spacing, and alignment to paragraphs."""
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
    fmt = paragraph.paragraph_format
    fmt.line_spacing = spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if alignment:
        paragraph.alignment = alignment


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = OxmlElement('w:r')
    
    # Set the run's text
    rPr = OxmlElement('w:rPr')
    
    # Apply styling (blue color and underline)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    
    # Set color to blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # Join all the xml elements together
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink


def add_contact_with_links(paragraph, contact_text):
    """Parse contact text and add hyperlinks for emails and URLs."""
    # Regex patterns
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    url_pattern = r'(?:https?://)?(?:www\.)?([a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}(?:/[^\s]*)?'
    
    # Find all emails and URLs with their positions
    items = []
    
    for match in re.finditer(email_pattern, contact_text):
        items.append({
            'start': match.start(),
            'end': match.end(),
            'text': match.group(),
            'type': 'email',
            'url': f'mailto:{match.group()}'
        })
    
    for match in re.finditer(url_pattern, contact_text):
        # Skip if this URL is already part of an email
        is_email = False
        for item in items:
            if item['type'] == 'email' and match.start() >= item['start'] and match.end() <= item['end']:
                is_email = True
                break
        
        if not is_email:
            url = match.group()
            # Add https:// if not present
            full_url = url if url.startswith('http') else f'https://{url}'
            items.append({
                'start': match.start(),
                'end': match.end(),
                'text': url,
                'type': 'url',
                'url': full_url
            })
    
    # Sort by position
    items.sort(key=lambda x: x['start'])
    
    # Build paragraph with mixed text and hyperlinks
    last_pos = 0
    for item in items:
        # Add plain text before the link
        if item['start'] > last_pos:
            run = paragraph.add_run(contact_text[last_pos:item['start']])
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        
        # Add the hyperlink
        add_hyperlink(paragraph, item['text'], item['url'])
        
        last_pos = item['end']
    
    # Add remaining text
    if last_pos < len(contact_text):
        run = paragraph.add_run(contact_text[last_pos:])
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)


def add_structured_contact(paragraph, contact_data):
    """
    Add structured contact information with proper hyperlinks.
    
    Expected contact_data format:
    {
        "phone": "+1 999 999 9999",
        "email": "abc@gmail.com",
        "linkedin": "https://linkedin.com/in/username",
        "github": "https://github.com/username",
        "portfolio": "https://myportfolio.com",
        ... any other links
    }
    """
    # Define the order of contact elements
    ordered_keys = ["phone", "email", "linkedin", "github", "portfolio", "website"]
    
    # Collect all items in order
    items = []
    
    # Add ordered items first
    for key in ordered_keys:
        if key in contact_data and contact_data[key]:
            items.append((key, contact_data[key]))
    
    # Add any additional links not in the predefined order
    for key, value in contact_data.items():
        if key not in ordered_keys and value:
            items.append((key, value))
    
    # Build the contact line
    for idx, (key, value) in enumerate(items):
        # Add separator before each item (except the first)
        if idx > 0:
            run = paragraph.add_run(" | ")
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        
        # Handle phone and email as plain text (email gets mailto link)
        if key == "phone":
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
        elif key == "email":
            add_hyperlink(paragraph, value, f"mailto:{value}")
        else:
            # For all other links (linkedin, github, portfolio, etc.)
            # Display the key name as a hyperlink
            display_text = key.capitalize()
            
            # Ensure URL has https:// prefix
            url = value
            if not url.startswith(('http://', 'https://')):
                url = f"https://{url}"
            
            add_hyperlink(paragraph, display_text, url)


def create_resume(data, file_name):
    doc = Document()

    # --- Name ---
    name = doc.add_paragraph()
    run = name.add_run(data["name"])
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    set_paragraph_format(name, font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # --- Contact Info ---
    contact = doc.add_paragraph()
    contact_data = data.get("contact", {})
    
    # Handle both old format (string) and new format (dict)
    if isinstance(contact_data, str):
        # Old format: single string with links auto-detected
        add_contact_with_links(contact, contact_data)
    else:
        # New format: structured object with phone, email, and links
        add_structured_contact(contact, contact_data)
    
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.line_spacing = 1.0
    contact.paragraph_format.space_before = Pt(0)
    contact.paragraph_format.space_after = Pt(0)

    # --- Summary ---
    heading = doc.add_heading("SUMMARY", level=1)
    set_paragraph_format(heading, font_size=11)
    heading.paragraph_format.space_before = Pt(0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set to black
    
    # Add horizontal line directly to the heading
    pPr = heading._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para = doc.add_paragraph(data["summary"])
    set_paragraph_format(para, font_size=10, spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # --- Technical Skills ---
    if data.get("technical_skills"):  # Add safety check
        heading = doc.add_heading("TECHNICAL SKILLS", level=1)
        set_paragraph_format(heading, font_size=11)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set to black
        
        # Add horizontal line directly to the heading
        pPr = heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for key, value in data["technical_skills"].items():
            skill_para = doc.add_paragraph()
            run = skill_para.add_run(f"{key}: ")
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            # Handle different value types: list, dict, or string
            if isinstance(value, list):
                value_text = ", ".join(value)
            elif isinstance(value, dict):
                # Handle nested dicts (like grouped subcategories)
                sub_parts = []
                for sub_key, sub_vals in value.items():
                    if isinstance(sub_vals, list):
                        sub_parts.append(f"{sub_key}: {', '.join(sub_vals)}")
                    else:
                        sub_parts.append(f"{sub_key}: {sub_vals}")
                value_text = "; ".join(sub_parts)
            else:
                value_text = str(value)
            
            skill_para.add_run(value_text)
            set_paragraph_format(skill_para, font_size=10, spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # --- Experience ---
    heading = doc.add_heading("EXPERIENCE", level=1)
    set_paragraph_format(heading, font_size=11)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set to black
    
    # Add horizontal line directly to the heading
    pPr = heading._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    for exp in data["experience"]:
        company_para = doc.add_paragraph()
        # Handle both 'role' and 'title' field names
        job_title = exp.get('role') or exp.get('title', 'Position')
        # Handle both 'period' and 'dates' field names
        dates = exp.get('period') or exp.get('dates', '')
        run = company_para.add_run(f"{exp['company']} | {job_title} | {dates}")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        set_paragraph_format(company_para, font_size=10, spacing=1.0)

        # Bullets aligned full width
        # Handle both 'points' and 'bullets' field names
        bullets = exp.get('points') or exp.get('bullets', [])
        for point in bullets:
            bullet = doc.add_paragraph(point, style="List Bullet")
            set_paragraph_format(bullet, font_size=10, spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # --- Projects ---
    if data.get("projects") and len(data["projects"]) > 0:
        heading = doc.add_heading("PROJECTS", level=1)
        set_paragraph_format(heading, font_size=11)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set to black
        
        # Add horizontal line directly to the heading
        pPr = heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for proj in data["projects"]:
            if proj.get("title"):  # Only add if title exists
                proj_title = doc.add_paragraph()
                run = proj_title.add_run(proj["title"])
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                set_paragraph_format(proj_title, font_size=10, spacing=1.0)
                
                # Project bullets
                for bullet in proj.get("bullets", []):
                    if bullet.strip():  # Only add non-empty bullets
                        bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                        set_paragraph_format(bullet_para, font_size=10, spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # --- Education ---
    heading = doc.add_heading("EDUCATION", level=1)
    set_paragraph_format(heading, font_size=11)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set to black
    
    # Add horizontal line directly to the heading
    pPr = heading._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    for edu in data["education"]:
        edu_para = doc.add_paragraph()
        run = edu_para.add_run(f"{edu['degree']}, {edu['institution']} ({edu['year']})")
        run.bold = True
        set_paragraph_format(edu_para, font_size=10, spacing=1.0)

    # --- Certifications ---
    if data.get("certifications") and len(data["certifications"]) > 0:
        heading = doc.add_heading("CERTIFICATIONS", level=1)
        set_paragraph_format(heading, font_size=11)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Set to black
        
        # Add horizontal line directly to the heading
        pPr = heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        for cert in data["certifications"]:
            if cert.get("name"):  # Only add if name exists
                cert_para = doc.add_paragraph()
                cert_text = cert["name"]
                if cert.get("organization"):
                    cert_text += f", {cert['organization']}"
                if cert.get("year"):
                    cert_text += f" ({cert['year']})"
                run = cert_para.add_run(cert_text)
                run.bold = True
                set_paragraph_format(cert_para, font_size=10, spacing=1.0)

    # --- Margins ---
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    folder_path = "./generated_resumes/"
    import os
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    file_name = os.path.join(folder_path, file_name)
    
    doc.save(file_name)
    print(f"✅ Document created successfully: {file_name}")

# Example usage with OLD format (backward compatible)
resume_data_old = {
    "name": "John Doe",
    "contact": "john.doe@gmail.com | 123-456-7890 | City, State | www.johndoe.com",
    "summary": "Experienced software developer with a strong background in Python and web development. Skilled in building scalable applications and working in agile teams.",
    "technical_skills": {
        "Programming Languages": ["Python", "JavaScript", "Java"],
        "Frameworks": ["Django", "React", "Flask"],
        "Databases": ["PostgreSQL", "MongoDB", "MySQL"],
        "Tools": {
            "Version Control": ["Git", "GitHub"],
            "CI/CD": ["Jenkins", "GitLab CI"]
        }
    },
    "experience": [
        {
            "company": "Tech Solutions Inc.",
            "title": "Senior Software Developer",
            "dates": "Jan 2020 - Present",
            "bullets": [
                "Led a team of 5 developers to build a scalable web application using Django and React.",
                "Implemented RESTful APIs and integrated third-party services to enhance application functionality.",
                "Optimized database queries, resulting in a 30% performance improvement."
            ]
        },
        {
            "company": "Web Innovations LLC",
            "title": "Software Developer",
            "dates": "Jun 2017 - Dec 2019",
            "bullets": [
                "Developed and maintained web applications using Flask and JavaScript.",
                "Collaborated with cross-functional teams to define project requirements and deliverables.",                    
                "Wrote unit tests and conducted code reviews to ensure code quality."
            ]
        }
    ],
    "education": [
        {
            "degree": "B.Sc. in Computer Science",
            "institution": "State University",
            "year": "2017"
        }
    ]
}

# Example usage with NEW structured contact format
resume_data_new = {
    "name": "Jane Smith",
    "contact": {
        "phone": "+1 773 359 3056",
        "email": "jane.smith@gmail.com",
        "linkedin": "https://linkedin.com/in/janesmith",
        "github": "https://github.com/janesmith",
        "portfolio": "https://janesmith.dev"
    },
    "summary": "Experienced software developer with a strong background in Python and web development. Skilled in building scalable applications and working in agile teams.",
    "technical_skills": {
        "Programming Languages": ["Python", "JavaScript", "Java"],
        "Frameworks": ["Django", "React", "Flask"],
        "Databases": ["PostgreSQL", "MongoDB", "MySQL"],
        "Tools": {
            "Version Control": ["Git", "GitHub"],
            "CI/CD": ["Jenkins", "GitLab CI"]
        }
    },
    "experience": [
        {
            "company": "Tech Solutions Inc.",
            "title": "Senior Software Developer",
            "dates": "Jan 2020 - Present",
            "bullets": [
                "Led a team of 5 developers to build a scalable web application using Django and React.",
                "Implemented RESTful APIs and integrated third-party services to enhance application functionality.",
                "Optimized database queries, resulting in a 30% performance improvement."
            ]
        }
    ],
    "education": [
        {
            "degree": "B.Sc. in Computer Science",
            "institution": "State University",
            "year": "2017"
        }
    ]
}

if __name__ == "__main__":
    # Test with new structured format
    create_resume(resume_data_new, "Jane_Smith_Resume_New_Format.docx")
    print("\n" + "="*50)
    # Test with old format (backward compatible)
    create_resume(resume_data_old, "John_Doe_Resume_Old_Format.docx")